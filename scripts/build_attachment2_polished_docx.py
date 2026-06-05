from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / ".omc" / "logs" / "attachment2-weekly-final-last-text.txt"
OUT = ROOT / "outputs" / "attachment2-weekly-final-polished.docx"
LABEL_W = 2200
VALUE_W = 6600


DATE_MAP = {
    "2025.9.15": "2025.9.16",
    "2025.9.22": "2025.9.23",
    "2025.9.29": "2025.9.30",
    "2025.10.13": "2025.10.14",
    "2025.10.20": "2025.10.21",
    "2025.10.27": "2025.10.29",
    "2025.11.10": "2025.11.11",
    "2025.11.24": "2025.11.25",
    "2025.12.8": "2025.12.9",
    "2025.12.22": "2025.12.23",
    "2026.1.5": "2026.1.6",
    "2026.1.19": "2026.1.20",
    "2026.2.2": "2026.2.4",
    "2026.2.20": "2026.2.21",
    "2026.3.2": "2026.3.4",
    "2026.3.16": "2026.3.18",
    "2026.3.30": "2026.3.31",
    "2026.4.13": "2026.4.15",
    "2026.4.27": "2026.4.28",
    "2026.5.11": "2026.5.12",
    "2026.5.29": "2026.5.28",
}


LABELS = [
    "指导日期",
    "指导地点",
    "参加人",
    "工作进展",
    "存在问题",
    "下一步工作安排",
    "指导教师(组)意见",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def para(doc: Document, text: str = "", style: str | None = None, align=None, first_indent: bool = False, size: float = 10.5, bold: bool = False):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_indent:
        pf.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 12, bold=True)


def set_table_borders(table, color="7F7F7F", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def cell_text(cell, text: str, bold: bool = False, align=None, size: float = 10.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths=(LABEL_W, VALUE_W)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="BFBFBF", size="4")
    for key, value in rows:
        row = table.add_row()
        set_cell_width(row.cells[0], widths[0])
        set_cell_width(row.cells[1], widths[1])
        set_cell_shading(row.cells[0], "F2F2F2")
        cell_text(row.cells[0], key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], value)
    para(doc, "")


def normalize_line(line: str) -> str:
    return line.replace("\u3000", "").strip()


def date_to_chinese(date_str: str) -> str:
    y, m, d = date_str.split(".")
    return f"{y}  年  {int(m)}  月  {int(d)}  日"


def parse_weekly_logs(lines: list[str]) -> list[dict[str, str]]:
    start = next(i for i, line in enumerate(lines) if normalize_line(line).startswith("第") and "周" in line)
    end = next(i for i in range(start, len(lines)) if normalize_line(lines[i]).startswith("题 目："))
    block_lines = [normalize_line(x) for x in lines[start:end]]

    entries: list[dict[str, str]] = []
    i = 0
    while i < len(block_lines):
        if not (block_lines[i].startswith("第") and "周" in block_lines[i]):
            i += 1
            continue
        entry: dict[str, str] = {"周次": block_lines[i]}
        i += 1
        current = None
        bucket: list[str] = []
        while i < len(block_lines) and not (block_lines[i].startswith("第") and "周" in block_lines[i]):
            line = block_lines[i]
            if line in LABELS:
                if current is not None:
                    entry[current] = "\n".join(x for x in bucket if x)
                current = line
                bucket = []
            elif line.startswith("指导教师（签字）"):
                if current is not None:
                    entry[current] = "\n".join(x for x in bucket if x)
                    current = None
                match = re.search(r"(\d{4})\s+年\s+(\d{1,2})\s+月\s+(\d{1,2})\s+日", line)
                if match:
                    entry["签字日期"] = f"{match.group(1)}.{int(match.group(2))}.{int(match.group(3))}"
            else:
                bucket.append(line)
            i += 1
        if current is not None:
            entry[current] = "\n".join(x for x in bucket if x)
        raw_date = entry.get("指导日期", "")
        entry["指导日期"] = DATE_MAP.get(raw_date, raw_date)
        entry["签字日期"] = DATE_MAP.get(entry.get("签字日期", raw_date), entry.get("签字日期", raw_date))
        entries.append(entry)
    return entries


def add_weekly_log(doc: Document, entry: dict[str, str]) -> None:
    add_heading(doc, entry["周次"], level=2)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="8C8C8C", size="6")

    rows = [
        ("指导日期", entry.get("指导日期", "")),
        ("指导地点", entry.get("指导地点", "")),
        ("参加人", entry.get("参加人", "")),
        ("工作进展", entry.get("工作进展", "")),
        ("存在问题", entry.get("存在问题", "")),
        ("下一步工作安排", entry.get("下一步工作安排", "")),
        ("指导教师(组)意见", entry.get("指导教师(组)意见", "")),
    ]
    for key, value in rows:
        row = table.add_row()
        set_cell_width(row.cells[0], LABEL_W)
        set_cell_width(row.cells[1], VALUE_W)
        set_cell_shading(row.cells[0], "F2F2F2")
        cell_text(row.cells[0], key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], value.replace("\n", "\n"))

    sig_row = table.add_row()
    set_cell_width(sig_row.cells[0], LABEL_W)
    set_cell_width(sig_row.cells[1], VALUE_W)
    cell_text(sig_row.cells[0], "指导教师（签字）：", align=WD_ALIGN_PARAGRAPH.LEFT)
    sig_date = entry.get("签字日期") or entry.get("指导日期", "")
    cell_text(sig_row.cells[1], f"指导日期： {date_to_chinese(sig_date)}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    para(doc, "")


def clean_english_text(text: str) -> str:
    fixes = {
        "se quentially": "sequentially",
        "seam lessly": "seamlessly",
        "perfor mance": "performance",
        "perfor mances": "performances",
        "en hance": "enhance",
        "represen tation": "representation",
        "connec tion": "connection",
        "parame ters": "parameters",
        "con volution": "convolution",
        "com putation": "computation",
        "descrip tor": "descriptor",
        "atten tion": "attention",
        "cardi nality": "cardinality",
        "de tection": "detection",
        "classi fication": "classification",
        "convolu tion": "convolution",
        "ap plicability": "applicability",
        "stack ing": "stacking",
        "FAR": "CIFAR",
        "CI CIFAR": "CIFAR",
    }
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace("con- firmed", "confirmed")
    for old, new in fixes.items():
        text = text.replace(old, new)
    return text


def clean_chinese_text(text: str) -> str:
    text = re.sub(r"\s+", "", text).strip()
    text = text.replace("tion", "")
    text = text.replace("......", "")
    text = text.replace("。。", "。")
    return text


def reflow_chunk(lines: list[str], language: str) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    buf: list[str] = []

    def flush(kind="p"):
        nonlocal buf
        if not buf:
            return
        text = " ".join(buf) if language == "en" else "".join(buf)
        text = clean_english_text(text) if language == "en" else clean_chinese_text(text)
        if text:
            out.append((kind, text))
        buf = []

    for raw in lines:
        line = normalize_line(raw)
        if not line:
            flush()
            continue
        if language == "en" and line == "tion" and out and out[-1][1].endswith("convolu"):
            kind, previous = out[-1]
            out[-1] = (kind, previous[:-7] + "convolution")
            continue
        if line.startswith("北京交通大学毕业设计"):
            flush()
            break
        if line in {"标题与译文格式与论文格式要求相同。插图内文字及图名也译成中文"}:
            continue
        is_heading = bool(re.match(r"^\d+(\.\d+)?\s+[\w\u4e00-\u9fff]", line)) or line.startswith(("Abstract", "Keywords", "摘要", "关键词", "贡献", "Contribution"))
        is_caption = line.startswith(("Fig.", "Table", "图 ", "表 "))
        if is_heading or is_caption:
            flush()
            out.append(("heading" if is_heading else "caption", clean_english_text(line) if language == "en" else clean_chinese_text(line)))
            continue
        buf.append(line)
        if language == "zh" and line.endswith(("。", "：", "；")):
            flush()
    flush()
    return out


def add_reflowed_section(doc: Document, items: list[tuple[str, str]], language: str) -> None:
    for kind, text in items:
        if kind == "heading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            set_run_font(r, size=12, bold=True)
        elif kind == "caption":
            p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
        else:
            p = para(doc, first_indent=(language == "zh"), size=10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(text)
            set_run_font(r, east_asia="宋体", ascii_font="Times New Roman", size=10.5)


def add_center_line(doc: Document, text: str, size: float = 10.5, bold: bool = False) -> None:
    para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, bold=bold)


def add_literature_header(doc: Document, language: str) -> None:
    if language == "en":
        add_center_line(doc, "CBAM: Convolutional Block Attention Module", size=13, bold=True)
        add_center_line(doc, "Sanghyun Woo, Jongchan Park, Joon-Young Lee, In So Kweon", size=10.5)
        add_key_value_table(
            doc,
            [
                ("1", "Korea Advanced Institute of Science and Technology, Daejeon, Korea"),
                ("2", "Lunit Inc., Seoul, Korea"),
                ("3", "Adobe Research, San Jose, CA, USA"),
                ("Email", "{shwoo93, iskweon77}@kaist.ac.kr; jcpark@lunit.io; jolee@adobe.com"),
            ],
            widths=(1200, 7600),
        )
    else:
        add_center_line(doc, "CBAM：卷积块注意力模块", size=13, bold=True)
        add_center_line(doc, "禹相铉、朴钟灿、李俊英、权仁昭", size=10.5)
        add_key_value_table(
            doc,
            [
                ("1", "韩国科学技术院，韩国大田"),
                ("2", "Lunit 公司，韩国首尔"),
                ("3", "Adobe 研究院，美国加利福尼亚州圣何塞"),
                ("邮箱", "{shwoo93, iskweon77}@kaist.ac.kr; jcpark@lunit.io; jolee@adobe.com"),
            ],
            widths=(1200, 7600),
        )


def add_cover(doc: Document) -> None:
    para(doc, "", size=14)
    para(doc, "", size=14)
    para(doc, "基于深度学习的交通零部件缺陷检测系统设计与实现", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    para(doc, "Design and Implementation of a Traffic Component Defect Detection System Based on Deep Learning", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    para(doc, "", size=12)
    rows = [
        ("学    院", "软件学院"),
        ("专    业", "软件工程"),
        ("学生姓名", "梁曦霖"),
        ("学    号", "22301094"),
        ("指导教师", "吴睿智"),
    ]
    for key, value in rows:
        para(doc, f"{key}：    {value}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    para(doc, "", size=12)
    para(doc, "北京交通大学", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    para(doc, "2026年6月", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()
    add_heading(doc, "工作日志列表", level=1)
    for item in [
        "1. 北京交通大学毕业设计（论文）任务书",
        "2. 北京交通大学毕业设计（论文）开题报告",
        "3. 北京交通大学毕业设计（论文）指导手册",
        "4. 北京交通大学毕业设计（论文）外文原文与译文",
    ]:
        para(doc, item, size=12)
    doc.add_page_break()


def add_pre_weekly_content(doc: Document, lines: list[str]) -> None:
    start = next(i for i, line in enumerate(lines) if normalize_line(line).startswith("题 目："))
    end = next(i for i, line in enumerate(lines) if normalize_line(line).startswith("题    目："))
    content = [normalize_line(x) for x in lines[start:end]]
    for line in content:
        if not line:
            continue
        if any(line.startswith(prefix) for prefix in ["题 目：", "毕业设计（论文）基本内容", "毕业设计（论文）拟解决", "毕业设计（论文）应完成", "参考资料推荐", "一、项目背景", "二、项目内容", "三、拟采取", "主要参考文献", "毕业设计（论文）进度安排", "指导教师意见"]):
            add_heading(doc, line, level=2)
        else:
            para(doc, line, first_indent=line.startswith(("本", "（", "论文", "为", "随", "近", "尽", "该")), size=10.5)
    doc.add_page_break()


def add_translation(doc: Document, lines: list[str]) -> None:
    start = next(i for i, line in enumerate(lines) if normalize_line(line).startswith("题 目：") and i > 600)
    trans = lines[start:]
    original_start = next(i for i, line in enumerate(trans) if normalize_line(line).startswith("CBAM:"))
    zh_start = next(i for i, line in enumerate(trans) if normalize_line(line).startswith("二、译文"))

    add_heading(doc, "北京交通大学毕业设计（论文）外文原文与译文", level=1)
    add_key_value_table(
        doc,
        [
            ("题目", "基于深度学习的交通零部件缺陷检测系统设计与实现"),
            ("学院", "软件学院"),
            ("专业", "软件工程"),
            ("学生姓名", "梁曦霖"),
            ("学号", "22301094"),
            ("文献来源", "CBAM: Convolutional Block Attention Module"),
        ],
    )
    add_heading(doc, "一、外文原文", level=1)
    add_literature_header(doc, "en")
    abstract_start = next(i for i in range(original_start, zh_start) if normalize_line(trans[i]).startswith("Abstract"))
    original_items = reflow_chunk(trans[abstract_start:zh_start], "en")
    add_reflowed_section(doc, original_items, "en")
    doc.add_page_break()

    add_heading(doc, "二、中文译文", level=1)
    add_literature_header(doc, "zh")
    zh_abs_start = next(i for i in range(zh_start, len(trans)) if normalize_line(trans[i]).startswith("摘要"))
    zh_items = reflow_chunk(trans[zh_abs_start:], "zh")
    add_reflowed_section(doc, zh_items, "zh")


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    add_cover(doc)
    add_pre_weekly_content(doc, lines)

    add_heading(doc, "北京交通大学毕业设计（论文）指导手册", level=1)
    add_key_value_table(
        doc,
        [
            ("题目", "基于深度学习的交通零部件缺陷检测系统设计与实现"),
            ("英文题目", "Design and Implementation of a Traffic Component Defect Detection System Based on Deep Learning"),
            ("学院", "软件学院"),
            ("专业", "软件工程"),
            ("学生姓名", "梁曦霖"),
            ("学生学号", "22301094"),
            ("指导教师", "吴睿智"),
            ("起止年月", "2025年9月-2026年6月"),
        ],
    )
    for note in [
        "1. 本手册全程记录学生开展本科毕业设计（论文）和教师指导的情况，由学生摘要填写，并由指导教师签字认可。",
        "2. 本科毕业设计（论文）开展过程中，每两周至少填写本手册1次。",
        "3. 本手册是检查学生本科毕业设计（论文）工作进展的重要方式，也是指导教师工作量认定的参考依据。",
        "4. 本手册是毕业设计（论文）工作日志的重要组成部分，在答辩时呈交答辩委员会，作为成绩评定的重要依据。",
        "5. 附页不够时可自行复印加页。",
    ]:
        para(doc, note, first_indent=True)
    para(doc, "北 京 交 通 大 学", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    para(doc, "2025年9月", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    for idx, entry in enumerate(parse_weekly_logs(lines)):
        add_weekly_log(doc, entry)
        if idx != 0 and idx % 3 == 2:
            doc.add_page_break()

    doc.add_page_break()
    add_translation(doc, lines)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
